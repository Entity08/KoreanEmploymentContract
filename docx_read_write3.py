import docx
from docx.shared import Pt, RGBColor

from law_dict import law_dict

def add_comment(doc, comment):  # 주석
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"({comment})")
    run.font.color.rgb = RGBColor(0, 102, 204)  # 주석 blue

def adjust_font_size(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size and run.font.size.pt < 11:  # 글씨 사이즈 11pt 이하인 경우 처리
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 224, 7)  # 처리한 글씨 색 green

def highlight_word(doc):
    new_word = []  # To track words for which comments have been added
    for para in doc.paragraphs:
        for run in para.runs:
            words = run.text.split()  # Split into words to check each one
            run.text = ''
            for word in words:
                new_run = para.add_run(word + ' ')
                if word in law_dict:
                    new_run.font.color.rgb = RGBColor(255, 0, 0)  # 하이라이트 red
                    if word not in new_word:
                        add_comment(doc, f"{word}: {law_dict[word]}")  # 주석 단어별로 1회 등장하도록 처리
                        new_word.append(word)
                else:
                    new_run.font.color.rgb = run.font.color.rgb  # Preserve original text color

def main():
    path = input('Enter the document path: ')
    doc = docx.Document(path)

    while True:
        print("---------------------")
        print("1. Adjust Font Size")
        print("2. Highlight")
        print("3. Save and Exit")
        print("---------------------")
        choice = input('Choose the menu: ')
        
        if choice == '2':
            highlight_word(doc)
        elif choice == '1':
            adjust_font_size(doc)
        elif choice == '3':
            new_doc_name = input('Enter new document name: ') + '.docx'
            doc.save(new_doc_name)
            break
        else:
            print("Invalid option, please try again.")

if __name__ == '__main__':
    main()
